from __future__ import annotations

import hashlib
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

#: Extensions we know how to read. Used to filter directory scans.
SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv", ".xlsx", ".xls"}


class UnsupportedFileType(ValueError):
    """Raised for a file extension no loader handles."""


def load_text_file(path: Path) -> str:
    """Read .txt / .md.

    errors="ignore" because real-world text files carry stray bytes from
    mixed encodings, and losing one character beats aborting the file.
    """
    return path.read_text(encoding="utf-8", errors="ignore")


def load_pdf(path: Path) -> str:
    """Extract text from a PDF, page by page.

    Two things to know:

    1. Pages are separated by a "[page N]" marker. That survives into
       the chunks, so an answer can cite the page it came from - much
       more useful than citing only the filename.
    2. This reads the *text layer*. A scanned PDF is images of text and
       has no text layer, so this returns "" - it isn't broken, it just
       needs OCR (a possible later addition).
    """
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[str] = []
    for number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"[page {number}]\n{text}")

    if not pages:
        logger.warning(
            "No text layer found in %s - it may be a scanned PDF needing OCR.",
            path.name,
        )
    return "\n\n".join(pages)


def load_docx(path: Path) -> str:
    """Extract text from a Word .docx.

    Covers paragraphs and tables. Tables get flattened to pipe-separated
    rows, which reads well enough for an LLM and keeps row grouping
    intact. Headers, footers and footnotes are not extracted.
    """
    import docx

    document = docx.Document(str(path))
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def load_tabular(path: Path) -> str:
    """Turn a .csv/.xlsx into text for RAG.

    IMPORTANT DISTINCTION: this is *not* how the analyst handles
    spreadsheets. The analyst loads them into DuckDB and runs real SQL,
    which is the right tool for "what's the median revenue by region?".

    This loader is for the different question "what does this file
    contain?" - it emits the schema plus a sample of rows so the RAG
    system can answer about a dataset's shape and columns. Embedding
    100k rows of numbers would be slow, expensive, and near-useless for
    semantic search.
    """
    import pandas as pd

    if path.suffix.lower() == ".csv":
        frame = pd.read_csv(path)
    else:
        frame = pd.read_excel(path)

    max_rows = 200
    header = (
        f"Table: {path.name}\n"
        f"Rows: {len(frame)}  Columns: {len(frame.columns)}\n"
        f"Column names: {', '.join(map(str, frame.columns))}\n"
    )
    if len(frame) > max_rows:
        header += f"(showing the first {max_rows} rows)\n"

    return header + "\n" + frame.head(max_rows).to_string(index=False)


#: Extension -> loader function. Adding a format means adding one
#: function above and one line here; nothing downstream changes.
LOADERS = {
    ".txt": load_text_file,
    ".md": load_text_file,
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".csv": load_tabular,
    ".xlsx": load_tabular,
    ".xls": load_tabular,
}


def load_document(path: Path) -> str:
    """Load any supported file to plain text.

    Raises:
        FileNotFoundError: path doesn't exist.
        UnsupportedFileType: no loader for this extension.
    """
    if not path.is_file():
        raise FileNotFoundError(f"Not a file: {path}")

    loader = LOADERS.get(path.suffix.lower())
    if loader is None:
        raise UnsupportedFileType(
            f"No loader for '{path.suffix}'. Supported: "
            f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    return loader(path)


def discover_files(root: Path, recursive: bool = True) -> list[Path]:
    """List loadable files under a path.

    A file path returns just that file (if supported); a directory
    returns every supported file inside it. Sorted so ingestion order is
    reproducible between runs.
    """
    if root.is_file():
        return [root] if root.suffix.lower() in SUPPORTED_EXTENSIONS else []

    pattern = "**/*" if recursive else "*"
    return sorted(
        p
        for p in root.glob(pattern)
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
    )


def file_hash(path: Path) -> str:
    """SHA-256 of a file's bytes, for skipping unchanged re-ingestion.

    Read in 64KB blocks so a large PDF never loads fully into memory
    just to be hashed.
    """
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(65536), b""):
            digest.update(block)
    return digest.hexdigest()
